#!/usr/bin/python

# Makes a summary of the object at the coordinates given on the command line,
# or if an input file is specified, the objects stored in the file.
#

import argparse
import getImg
import cv2
import os

from docx import Document
import docx.shared
import docx.enum.table

OUTDIR = "farmSummaries"
IMSIZE = 128
TEMPLATE_FNAME = "Survey_Record_Blank.docx"

def writeSummaryForm(fpath,name,lon,lat,imgList):
    """ Create a microsoft word form including information that we have 
    collected.
    """
    addrStr,parishStr = getImg.getAddress(lat,lon)

    templatePath = os.path.join(
        os.path.dirname(os.path.realpath(__file__)),
        TEMPLATE_FNAME)
    document = Document(templatePath)

    #print(dir(document))
    #print(document.tables)

    t = document.tables[0]

    # Delete the last row of the table (so we can just extend the image rows)
    # from https://groups.google.com/forum/#!topic/python-docx/aDumlNvK6GM
    tbl = t._tbl
    tr = t.rows[len(t.rows)-1]._tr
    tbl.remove(tr)


    # Add Coordinates
    rowNo = 1
    t.cell(rowNo,3).add_paragraph("(%f,%f)" % (lat,lon))
    osgb_n,osgb_e = getImg.latlon2osgb(lat,lon)
    t.cell(rowNo,3).add_paragraph("(%.0f,%.0f)" % (osgb_n,osgb_e))
    t.cell(rowNo,2).add_paragraph(parishStr)
    t.rows[rowNo].height_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.AUTO
    
    # Add Address and Farm Name
    rowNo = 2
    t.cell(rowNo,2).add_paragraph(addrStr)
    t.cell(rowNo,3).add_paragraph(name)
    t.rows[rowNo].height_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.AUTO

    rowNo = 5
    imgHt = 50 # mm
    for img in imgList:
        t.rows[rowNo].height_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.AUTO
        t.cell(rowNo,0).add_paragraph().add_run().add_picture(img['fpath'],
                                                              width=docx.shared.Mm(imgHt))
        #print getImg.tileSeries
        seriesStr = img['series']
        t.cell(rowNo,3).add_paragraph().text="%s - %s" % \
                                              (seriesStr,
                                               getImg.tileSeries[seriesStr]['dates'])
        rowNo = rowNo + 1
        if rowNo>7:
            r = t.add_row()
            r.cells[0].merge(r.cells[1])
            r.cells[2].merge(r.cells[3])

    
    document.save(fpath)


                        
def makeSummary(name,lon,lat):
    """ Make a summary of the object called name at coordinates (lon,lat)
    """
    imgList = []
    outDir = os.path.join(os.getcwd(),OUTDIR)
    if (not os.path.exists(outDir)):
        os.makedirs(outDir)

    for seriesStr in ["OS_6in_1st","OS_10k","OSM","Google","Google-Satellite"]:
        fpath = os.path.join(outDir,"%s_%s.png" % (name,seriesStr))
        img,bigImg,noFarmImg = getImg.getFarmImgLatLon(seriesStr,lat,lon,IMSIZE)
        cv2.imwrite(fpath,img)
        imgList.append({'series':seriesStr,'fpath':fpath})
    #print imgList

    summaryFpath = os.path.join(outDir,"Summary_%s.docx" % (name))
    writeSummaryForm(summaryFpath,name,lon,lat,imgList)

if (__name__=="__main__"):
    # construct the argument parser and parse the arguments
    ap = argparse.ArgumentParser()
    ap.add_argument("-f", "--infile", required=False,
                    help="input file, which should be CSV format name, lon, lat")
    ap.add_argument("-c", "--coords", required=False,
                    help="Coordinates of single object to summarise (lon,lat)")
    ap.add_argument("-n", "--name", required=False,
                    help="Name of object - used in filenames")
    
    args = vars(ap.parse_args())
    print args


    if (args["infile"]!=None):
        print("FIXME - Handle input file of farms to analyse")

    else:
        if (args["coords"]==None and args["name"]==None):
            print("ERROR - Both coords and name are required if no input file is specified")
        else:
            lat = float(args['coords'].strip("(").strip(")").split(",")[0])
            lon = float(args['coords'].strip("(").strip(")").split(",")[1])
            name = args['name']
            print("Name=%s, Coords = (%f,%f)" % (name,lat,lon))

        makeSummary(name,lon,lat)

        
